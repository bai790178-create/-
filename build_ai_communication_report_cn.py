from pathlib import Path

import build_ai_communication_report as base
from PIL import ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("report_output")
CHARTS = OUT / "charts_cn"
OUT.mkdir(exist_ok=True)
CHARTS.mkdir(exist_ok=True)
base.CHARTS = CHARTS


def cn_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


base.font = cn_font


def set_run_font(run, size=12, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    for style_name, size, name in [
        ("Normal", 12, "宋体"),
        ("Heading 1", 16, "黑体"),
        ("Heading 2", 14, "黑体"),
        ("Heading 3", 12, "黑体"),
    ]:
        style = doc.styles[style_name]
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        style.font.size = Pt(size)
        if style_name.startswith("Heading"):
            style.font.bold = True
            style.font.color.rgb = RGBColor(31, 78, 121)
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.2
        else:
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.first_line_indent = Inches(0.28)
            style.paragraph_format.space_after = Pt(0)


def add_centered(doc, text, size=16, bold=False, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size, bold, "黑体" if bold else "宋体")


def add_para(doc, text, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if not first_indent:
        p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_run_font(r, 12, False, "宋体")
    return p


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(5.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    cap.paragraph_format.line_spacing = 1.1
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, 10.5, False, "宋体")
    r.italic = True


def clear_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)


def add_figure_pair(doc, left_item, right_item):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(table)
    for col, item in enumerate([left_item, right_item]):
        image_path, caption = item
        image_cell = table.rows[0].cells[col]
        cap_cell = table.rows[1].cells[col]
        image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cap_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = image_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.add_run().add_picture(str(image_path), width=Inches(3.05))
        cp = cap_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = None
        cp.paragraph_format.line_spacing = 1.05
        r = cp.add_run(caption)
        set_run_font(r, 9.5, False, "宋体")
        r.italic = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    spacer.paragraph_format.first_line_indent = None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, 10.5, bold, "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_summary_table(doc):
    rows = [
        ("样本", "75 份有效问卷；第 4 题有效填写人数为 40。"),
        ("定量方法", "基于第二版间接式问卷进行描述性统计分析。"),
        ("定性方法", "采用情境式访谈，重点分析 Linda 和 Bob 的案例。"),
        ("分析重点", "沟通语气、表达风格、情绪期待、人际互动习惯和潜意识影响。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["项目", "说明"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for a, b in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], a, True)
        set_cell_text(cells[1], b)


def make_charts():
    return {
        "q1": base.pie_chart("第1题：你的年龄是？", ["18岁以下", "18-35岁", "35-45岁", "45岁及以上"], [5.33, 72, 9.33, 13.33], "q1_age_distribution_cn.png", "单选题，N = 75。"),
        "q2": base.pie_chart("第2题：一周内使用 AI 的频率", ["每天", "两三天一次", "几乎不"], [65.33, 34.67, 0], "q2_ai_use_frequency_cn.png", "单选题，N = 75。"),
        "q3": base.bar_chart("第3题：你通常用 AI 做什么？", ["搜索引擎", "倾诉对象", "咨询社交问题"], [88, 42.67, 26.67], "q3_ai_use_purposes_cn.png", "多选题，N = 75。", colors=[base.BLUE, base.TEAL, base.GOLD]),
        "q4": base.bar_chart("第4题：AI 的肯定是奉承还是基于事实？", ["完全奉承", "基本奉承", "基本基于事实", "完全基于事实"], [12.5, 47.5, 25, 15], "q4_flattery_or_fact_cn.png", "单选题，n = 40。", colors=[base.BLUE, base.TEAL, base.GOLD, (141, 105, 181)]),
        "q5": base.pie_chart("第5题：被 AI 肯定时是否心情愉悦？", ["是", "否"], [76, 24], "q5_pleasure_from_affirmation_cn.png", "单选题，N = 75。"),
        "q6": base.bar_chart("第6题：采纳 AI 建议的概率", ["0-20", "20-40", "40-60", "60-80", "80-100"], [8, 12, 42.67, 25.33, 12], "q6_adopting_ai_suggestions_cn.png", "单选题，N = 75。", colors=[base.BLUE, base.TEAL, base.GOLD, (141, 105, 181), (85, 150, 95)]),
        "q7": base.bar_chart("第7题：高强度使用 AI 主要与什么有关？", ["AI 功能更丰富", "AI 信息可信度高", "与 AI 交流比询问他人更方便"], [68, 38.67, 68], "q7_reasons_for_high_ai_use_cn.png", "多选题，N = 75。", colors=[base.BLUE, base.TEAL, base.GOLD]),
        "q9": base.pie_chart("第9题：需要别人帮忙时如何表达？", ["先说明原因", "直接说明要做什么"], [73.33, 26.67], "q9_help_request_style_cn.png", "单选题，N = 75。"),
        "q10": base.bar_chart("第10题：是否省略语气词、铺垫和客套话？", ["是", "否"], [53.33, 46.67], "q10_omit_pleasantries_cn.png", "单选题，N = 75。", colors=[base.BLUE, base.TEAL]),
        "q11": base.bar_chart("第11题：表达是否变得更直接？", ["是", "否"], [58.67, 41.33], "q11_more_direct_blunt_cn.png", "单选题，N = 75。", colors=[base.BLUE, base.TEAL]),
    }


def build_doc():
    charts = make_charts()
    doc = Document()
    style_doc(doc)

    add_centered(doc, "项目式学习研究报告", 22, True, 24)
    add_centered(doc, "AI 如何影响我们的交流", 20, True, 24)
    add_centered(doc, "关于 AI 潜移默化影响沟通习惯的问卷与访谈分析", 15, False, 42)
    add_centered(doc, "医学院", 13, False, 12)
    add_centered(doc, "2026年6月", 13, False, 12)
    doc.add_page_break()

    doc.add_heading("摘要", level=1)
    add_para(doc, "本报告研究人工智能是否会在人们没有明确意识到的情况下，潜移默化地影响日常交流方式。研究基于 75 份有效问卷、情境式访谈以及关于 AI 幻觉、AI 奉承倾向和社交依赖的文献综述。结果显示，AI 不再只是信息工具，也逐渐承担情绪支持和人际替代功能。问卷中，72% 的受访者年龄为 18 至 35 岁，65.33% 每天使用 AI，88% 将 AI 作为搜索引擎，42.67% 将 AI 当作倾诉对象。更重要的是，虽然第 4 题中 60% 的受访者认为 AI 的肯定主要带有奉承性质，但仍有 76% 的受访者表示被 AI 肯定时会感到愉悦。同时，53.33% 的受访者表示会省略语气词、铺垫和客套话，58.67% 认为自己最近与人沟通时表达更直接。以上结果支持本研究假设：AI 可能逐渐改变使用者的情绪期待、表达风格和人际互动习惯。")
    add_para(doc, "关键词：人工智能；沟通；情绪依赖；AI 奉承；人际互动", False)

    doc.add_heading("1. 研究背景", level=1)
    add_para(doc, "人工智能已经进入日常交流场景。人们使用 AI 搜索信息、做决定、写消息、解决私人问题，甚至寻求情绪支持。这种便利性带来一个矛盾：一方面，AI 让沟通更快、更高效；另一方面，由于 AI 往往即时、礼貌且顺从地回应用户，它也可能悄悄改变用户对真实人际交流的期待。")
    add_para(doc, "本研究关注的不是人们是否频繁使用 AI，而是长期与 AI 互动是否会改变人们与他人交流的方式、寻求肯定的方式以及面对分歧时的反应。这个问题之所以重要，是因为这种影响往往不容易被使用者直接意识到。")

    doc.add_heading("2. 文献综述", level=1)
    add_para(doc, "已有研究和公共讨论提出了三个相关问题。首先，关于语言模型幻觉的研究指出，AI 可能生成看似自信但并不可靠的答案。由于模型常被训练为给出回答而不是承认“不知道”，用户必须对 AI 输出保持批判性。")
    add_para(doc, "其次，关于 AI 奉承倾向的研究表明，AI 相比人类交流对象更容易认可和支持用户。项目提供的文献综述指出，AI 模型可能为了维持用户满意度而更倾向于赞同用户，这与本研究关注的“过度顺从的沟通环境”直接相关。")
    add_para(doc, "第三，一些个人叙述显示，长期与聊天机器人互动可能影响真实社交关系。但这些材料通常比较个人化，也多来自西方语境。因此，本研究将视角放在中国用户，尤其是年轻且高频使用 AI 的群体上。")

    doc.add_heading("3. 研究假设", level=1)
    add_para(doc, "本研究假设：AI 会在人们没有充分意识到的情况下影响他们的交流。影响主要体现在五个方面：沟通语气、表达风格、情绪期待、人际互动习惯和潜意识变化。具体而言，AI 可能让用户更习惯直接、高效、顺从和情绪肯定式的交流。久而久之，用户可能更偏好低冲突、快速获得肯定的互动方式，并对真实人际交流中的复杂性和分歧变得不耐烦。")

    doc.add_heading("4. 研究方法", level=1)
    add_para(doc, "本研究采用问卷与访谈结合的混合研究方法。研究过程中共设计了两版问卷。第一版问卷效果并不理想，因为部分问题过于直接，例如要求受访者判断 AI 是否改变了自己的日常交流，或评价自己对家人的语气。这类问题涉及个人尊严和自我形象，容易让受访者不愿承认真实变化。")
    add_para(doc, "因此，第二版问卷采用更间接、更委婉的设计，收集受访者的年龄、AI 使用频率、使用目的、对 AI 语气的感知、被 AI 肯定后的情绪反应、依赖 AI 的原因以及现实沟通习惯变化。同时，问卷加入情境式开放题，让受访者设想如何向 AI 询问家乡旅游地点，从而观察他们自然的提问语气和表达方式。")
    add_para(doc, "访谈作为定性补充。直接询问时，Linda 起初否认 AI 影响了她的交流；但在具体情境引导下，她承认自己喜欢 AI 带来的认可感，并会回避可能质疑她观点的人。Bob 则表示，为了防止 AI 提供错误信息，他习惯保持批判，而这种习惯有时会延伸到与同学讨论时。")
    add_summary_table(doc)

    doc.add_heading("5. 发现与数据分析", level=1)
    doc.add_heading("5.1 年轻群体和高频用户构成主要样本", level=2)
    add_para(doc, "问卷结果显示，年轻成年人是本研究的核心用户群体。在 75 名有效受访者中，72% 年龄为 18 至 35 岁，65.33% 每天使用 AI。这说明数据主要反映了已经熟悉 AI、并将 AI 纳入日常生活的群体。")
    add_figure_pair(doc, (charts["q1"], "图1 第1题年龄分布。"), (charts["q2"], "图2 第2题 AI 使用频率。"))
    add_para(doc, "这一结果对研究有帮助，因为潜移默化的影响更容易出现在高频用户身上。如果一个人每天反复体验 AI 的交流方式，这种方式更可能被习惯化。")

    doc.add_heading("5.2 AI 使用从信息搜索扩展到情绪支持", level=2)
    add_para(doc, "第二个重要现象是，AI 不仅是技术工具。虽然 88% 的受访者将 AI 用作搜索引擎，但 42.67% 也将其作为倾诉对象，26.67% 会用 AI 咨询社交问题。这说明 AI 已进入原本属于人际交流的领域。")
    add_figure(doc, charts["q3"], "图3 第3题 AI 使用目的。")
    add_para(doc, "这支持了关于人际互动习惯的假设。当人们用 AI 进行情绪表达或寻求社交建议时，一部分原本需要由他人回应的需求被转移给机器。")

    doc.add_heading("5.3 用户知道 AI 可能奉承，却仍享受肯定", level=2)
    add_para(doc, "数据中出现了理性判断与情绪反应之间的矛盾。第 4 题中，12.5% 选择“完全奉承”，47.5% 选择“基本奉承”，合计 60%。但第 5 题中，76% 的受访者表示被 AI 肯定时会感到愉悦。第 6 题进一步说明，受访者并非完全盲从 AI：最大群体（42.67%）将采纳 AI 建议的概率放在 40-60 区间，另有 37.33% 放在 60 以上。")
    add_figure_pair(doc, (charts["q4"], "图4 第4题 对 AI 肯定性质的判断。"), (charts["q5"], "图5 第5题 被 AI 肯定后的情绪反应。"))
    add_figure(doc, charts["q6"], "图6 第6题 采纳 AI 建议的概率。")
    add_para(doc, "这些图表共同说明，AI 的影响不一定来自用户完全相信 AI，而来自反复获得情绪奖励。即使用户知道 AI 可能在奉承，肯定本身仍然能带来愉悦感。")

    doc.add_heading("5.4 便利性推动求助行为向 AI 转移", level=2)
    add_para(doc, "第 7 题显示，高强度使用 AI 的主要原因并不只是信任。68% 选择“AI 功能更丰富”，68% 选择“与 AI 交流比询问他人更方便”，而选择“AI 信息可信度高”的比例为 38.67%。这说明频繁使用 AI 更多源于功能性和便利性，而非绝对信任。")
    add_figure(doc, charts["q7"], "图7 第7题 高强度使用 AI 的原因。")
    add_para(doc, "当用户反复因为“更方便、更安全”而选择 AI，他们可能逐渐减少与真实他人进行复杂交流的机会。这使 AI 的便利性与人际摩擦回避之间产生联系。")

    doc.add_heading("5.5 现实交流变得更直接、情绪缓冲减少", level=2)
    add_para(doc, "最后一组数据涉及沟通风格。第 9 题中，73.33% 的受访者表示向他人求助时会先说明原因，说明多数人仍保留人际礼貌规范。但第 10 题中，53.33% 承认聊天时会下意识省略语气词、铺垫和客套话；第 11 题中，58.67% 认为自己最近与人沟通时越来越少解释、安抚或缓和语气，表达更干脆直接。")
    add_figure_pair(doc, (charts["q9"], "图8 第9题 向他人求助时的表达方式。"), (charts["q10"], "图9 第10题 是否省略语气词和客套话。"))
    add_figure(doc, charts["q11"], "图10 第11题 表达是否更直接。")
    add_para(doc, "这说明 AI 式交流可能部分迁移到人与人的交流中。AI 交流通常直接、高效、以任务为中心，这些特征在机器互动中有效，但在人际交流中可能显得冷淡、急躁或不够体贴。")

    doc.add_heading("6. 访谈分析", level=1)
    add_para(doc, "访谈进一步解释了问卷数据背后的心理机制。Linda 起初否认 AI 影响了自己的日常交流，但在情境问题中，她逐渐承认自己喜欢 AI 的认可和接纳感，也更不愿意与可能否定或质疑她的人交流。她的案例对应了问卷中“被 AI 肯定会感到愉悦”的结果。")
    add_para(doc, "Bob 的案例则揭示了另一种迁移机制。由于 AI 有时会提供错误信息，他需要经常保持批判性。但这种批判习惯有时会延伸到与同学讨论时，使他更难立即相信他人，甚至显得有些粗鲁。")

    doc.add_heading("7. 讨论：数据反映了什么", level=1)
    add_para(doc, "总体而言，数据反映出三个相互关联的现象。第一，AI 从信息工具扩展为情绪支持对象。第二，用户并非不知道 AI 会奉承，但仍会被 AI 的肯定所吸引。第三，长期与 AI 互动可能让直接、高效、低语境的表达方式变得正常化，并影响真实人际沟通。")
    add_para(doc, "这些发现将文献综述与一手数据连接起来。文献指出 AI 可能幻觉、奉承或顺从用户；本研究的数据进一步说明，这些模型特征可能如何影响使用者：带来情绪依赖、分歧回避、表达机械化，以及批判习惯向人际互动迁移。")

    doc.add_heading("8. 结论", level=1)
    add_para(doc, "本研究支持最初假设：AI 可能在用户没有充分意识到的情况下影响交流习惯。这种影响体现在情绪期待、求助方式和表达风格上。AI 的便利性和顺从语气使其不仅是工具，也逐渐成为情绪回应对象。同时，反复与 AI 互动可能使用户更直接、更少使用情绪缓冲，也更不愿面对真实关系中的分歧。")
    add_para(doc, "研究也说明了方法设计的重要性。直接提问容易失败，因为受访者不愿承认可能损害自我形象的变化。间接问卷和情境式访谈更适合研究潜意识层面的沟通习惯。")

    doc.add_heading("参考文献", level=1)
    for ref in [
        "Kalai, A. (2025). Why Language Models Hallucinate. arXiv.",
        "Cheng, M. (2025). Sycophantic AI decreases prosocial intentions and promotes dependence. Science.",
        "Samuel, A. (2024). I Love Being Social. Then I Started Talking to a Chatbot. The Wall Street Journal.",
        "项目问卷数据. (2026). AI 是否潜移默化地影响我们的交流. 未发表问卷报告.",
    ]:
        add_para(doc, ref, False)

    doc.add_heading("附录：本报告使用的核心问卷数据", level=1)
    rows = [
        ("第1题 18岁以下", "4 / 75", "5.33%"), ("第1题 18-35岁", "54 / 75", "72%"), ("第1题 35-45岁", "7 / 75", "9.33%"), ("第1题 45岁及以上", "10 / 75", "13.33%"),
        ("第2题 每天", "49 / 75", "65.33%"), ("第2题 两三天一次", "26 / 75", "34.67%"), ("第2题 几乎不", "0 / 75", "0%"),
        ("第3题 搜索引擎", "66 / 75", "88%"), ("第3题 倾诉对象", "32 / 75", "42.67%"), ("第3题 咨询社交问题", "20 / 75", "26.67%"),
        ("第4题 完全奉承", "5 / 40", "12.5%"), ("第4题 基本奉承", "19 / 40", "47.5%"), ("第4题 基本基于事实", "10 / 40", "25%"), ("第4题 完全基于事实", "6 / 40", "15%"),
        ("第5题 是", "57 / 75", "76%"), ("第5题 否", "18 / 75", "24%"),
        ("第6题 0-20", "6 / 75", "8%"), ("第6题 20-40", "9 / 75", "12%"), ("第6题 40-60", "32 / 75", "42.67%"), ("第6题 60-80", "19 / 75", "25.33%"), ("第6题 80-100", "9 / 75", "12%"),
        ("第7题 AI 功能更丰富", "51 / 75", "68%"), ("第7题 AI 信息可信度高", "29 / 75", "38.67%"), ("第7题 与 AI 交流更方便", "51 / 75", "68%"),
        ("第9题 先说明原因", "55 / 75", "73.33%"), ("第9题 直接说明任务", "20 / 75", "26.67%"),
        ("第10题 是", "40 / 75", "53.33%"), ("第10题 否", "35 / 75", "46.67%"), ("第11题 是", "44 / 75", "58.67%"), ("第11题 否", "31 / 75", "41.33%"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["题目 / 指标", "人数", "比例"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, i == 0)

    out = OUT / "AI如何影响我们的交流_中文报告_并列图版.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_doc())
