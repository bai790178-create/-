from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("report_output")
CHARTS = OUT / "charts"
OUT.mkdir(exist_ok=True)
CHARTS.mkdir(exist_ok=True)


BLUE = (44, 116, 181)
TEAL = (58, 170, 165)
GOLD = (218, 165, 32)
GRAY = (112, 112, 112)
LIGHT_GRAY = (232, 236, 241)
DARK = (32, 32, 32)


def font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def text_width(draw, text, fnt):
    bbox = draw.textbbox((0, 0), text, font=fnt)
    return bbox[2] - bbox[0]


def bar_chart(title, labels, values, filename, subtitle=None, max_value=100, colors=None):
    width, height = 1400, 820
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42, True)
    sub_font = font(24)
    label_font = font(27)
    value_font = font(26, True)
    axis_font = font(22)

    draw.text((70, 48), title, fill=DARK, font=title_font)
    if subtitle:
        draw.text((70, 103), subtitle, fill=GRAY, font=sub_font)

    left, right, top, bottom = 420, 1280, 190, 720
    n = len(labels)
    gap = 24
    bar_h = int((bottom - top - gap * (n - 1)) / n)
    colors = colors or [BLUE] * n

    for tick in range(0, max_value + 1, 20):
        x = left + int((right - left) * tick / max_value)
        draw.line((x, top - 8, x, bottom + 8), fill=(235, 235, 235), width=2)
        draw.text((x - 14, bottom + 22), f"{tick}%", fill=GRAY, font=axis_font)

    for i, (label, value) in enumerate(zip(labels, values)):
        y = top + i * (bar_h + gap)
        draw.rounded_rectangle((left, y, right, y + bar_h), radius=14, fill=LIGHT_GRAY)
        bar_w = int((right - left) * value / max_value)
        draw.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=14, fill=colors[i % len(colors)])
        wrapped = label
        while text_width(draw, wrapped, label_font) > left - 110 and " " in wrapped:
            parts = wrapped.split(" ")
            mid = len(parts) // 2
            wrapped = " ".join(parts[:mid]) + "\n" + " ".join(parts[mid:])
            break
        lines = wrapped.split("\n")
        for j, line in enumerate(lines):
            draw.text((70, y + 8 + j * 31), line, fill=DARK, font=label_font)
        draw.text((left + bar_w + 18, y + bar_h / 2 - 16), f"{value:.2f}%" if value % 1 else f"{int(value)}%", fill=DARK, font=value_font)

    out = CHARTS / filename
    img.save(out, quality=95)
    return out


def grouped_bar_chart(title, groups, series, filename, subtitle=None):
    width, height = 1400, 780
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42, True)
    sub_font = font(24)
    label_font = font(25)
    value_font = font(22, True)
    axis_font = font(22)
    legend_font = font(24)

    draw.text((70, 45), title, fill=DARK, font=title_font)
    if subtitle:
        draw.text((70, 100), subtitle, fill=GRAY, font=sub_font)

    left, right, top, bottom = 430, 1280, 190, 640
    for tick in range(0, 101, 20):
        x = left + int((right - left) * tick / 100)
        draw.line((x, top - 8, x, bottom + 8), fill=(235, 235, 235), width=2)
        draw.text((x - 14, bottom + 24), f"{tick}%", fill=GRAY, font=axis_font)

    colors = [BLUE, TEAL, GOLD]
    row_h = (bottom - top) / len(groups)
    bar_h = 34
    for i, group in enumerate(groups):
        base_y = int(top + i * row_h + 18)
        draw.text((70, base_y + 18), group, fill=DARK, font=label_font)
        for j, s in enumerate(series):
            val = s["values"][i]
            y = base_y + j * (bar_h + 8)
            draw.rounded_rectangle((left, y, right, y + bar_h), radius=9, fill=LIGHT_GRAY)
            bar_w = int((right - left) * val / 100)
            draw.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=9, fill=colors[j])
            draw.text((left + bar_w + 12, y + 3), f"{val:.2f}%" if val % 1 else f"{int(val)}%", fill=DARK, font=value_font)

    legend_x = 70
    for j, s in enumerate(series):
        y = 690
        x = legend_x + j * 360
        draw.rounded_rectangle((x, y, x + 28, y + 28), radius=5, fill=colors[j])
        draw.text((x + 38, y - 1), s["name"], fill=DARK, font=legend_font)

    out = CHARTS / filename
    img.save(out, quality=95)
    return out


def pie_chart(title, labels, values, filename, subtitle=None):
    width, height = 1200, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(40, True)
    sub_font = font(24)
    label_font = font(27)
    draw.text((70, 45), title, fill=DARK, font=title_font)
    if subtitle:
        draw.text((70, 98), subtitle, fill=GRAY, font=sub_font)

    colors = [BLUE, TEAL, GOLD, (141, 105, 181)]
    cx, cy, r = 360, 430, 230
    start = -90
    total = sum(values)
    for i, val in enumerate(values):
        angle = val / total * 360
        draw.pieslice((cx - r, cy - r, cx + r, cy + r), start, start + angle, fill=colors[i], outline="white", width=4)
        start += angle

    y = 245
    for i, (label, val) in enumerate(zip(labels, values)):
        draw.rounded_rectangle((720, y, 750, y + 30), radius=5, fill=colors[i])
        draw.text((770, y - 2), f"{label}: {val:.2f}%" if val % 1 else f"{label}: {int(val)}%", fill=DARK, font=label_font)
        y += 68

    out = CHARTS / filename
    img.save(out, quality=95)
    return out


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.first_line_indent = Inches(0.28)
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        s = styles[name]
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(31, 78, 121)
        s.paragraph_format.first_line_indent = None
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.2


def add_centered(doc, text, size=16, bold=False, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.bold = bold


def add_para(doc, text, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    if not first_indent:
        p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    return p


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    cap.paragraph_format.line_spacing = 1.1
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(10.5)


def add_summary_table(doc):
    rows = [
        ("Sample", "75 valid questionnaire responses; Q4 has 40 valid responses."),
        ("Quantitative method", "Descriptive statistics from the second, indirect questionnaire."),
        ("Qualitative method", "Situational interviews with students, including Linda and Bob."),
        ("Analytical focus", "Tone, style, emotional expectation, interpersonal habits, and subconscious influence."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Item", True)
    set_cell_text(hdr[1], "Description", True)
    set_cell_shading(hdr[0], "F2F4F7")
    set_cell_shading(hdr[1], "F2F4F7")
    for a, b in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], a, True)
        set_cell_text(cells[1], b)


def make_charts():
    return {
        "q1": pie_chart(
            "Q1. What Is Your Age?",
            ["Under 18", "18-35", "35-45", "45 and above"],
            [5.33, 72, 9.33, 13.33],
            "q1_age_distribution.png",
            "Single-choice question, N = 75.",
        ),
        "q2": pie_chart(
            "Q2. How Often Do You Use AI in One Week?",
            ["Every day", "Once every two or three days", "Rarely"],
            [65.33, 34.67, 0],
            "q2_ai_use_frequency.png",
            "Single-choice question, N = 75.",
        ),
        "q3": bar_chart(
            "Q3. What Do You Usually Use AI For?",
            ["Search engine", "Confidant / emotional object", "Advice on social problems"],
            [88, 42.67, 26.67],
            "q3_ai_use_purposes.png",
            "Multiple-choice question, N = 75.",
            colors=[BLUE, TEAL, GOLD],
        ),
        "q4": bar_chart(
            "Q4. Is AI's Affirmation Flattery or Fact-Based?",
            ["Completely flattering", "Basically flattering", "Basically fact-based", "Completely fact-based"],
            [12.5, 47.5, 25, 15],
            "q4_flattery_or_fact.png",
            "Single-choice question, n = 40.",
            colors=[BLUE, TEAL, GOLD, (141, 105, 181)],
        ),
        "q5": pie_chart(
            "Q5. Do You Feel Pleased When Affirmed by AI?",
            ["Yes", "No"],
            [76, 24],
            "q5_pleasure_from_affirmation.png",
            "Single-choice question, N = 75.",
        ),
        "q6": bar_chart(
            "Q6. Probability of Adopting AI's Suggestions",
            ["0-20", "20-40", "40-60", "60-80", "80-100"],
            [8, 12, 42.67, 25.33, 12],
            "q6_adopting_ai_suggestions.png",
            "Single-choice question, N = 75.",
            colors=[BLUE, TEAL, GOLD, (141, 105, 181), (85, 150, 95)],
        ),
        "q7": bar_chart(
            "Q7. What Is High-Intensity AI Use Related To?",
            ["AI has richer functions", "AI information is more reliable", "AI is more convenient than asking others"],
            [68, 38.67, 68],
            "q7_reasons_for_high_ai_use.png",
            "Multiple-choice question, N = 75.",
            colors=[BLUE, TEAL, GOLD],
        ),
        "q9": pie_chart(
            "Q9. How Do You Ask Others for Help?",
            ["Explain the reason first", "Directly state the task"],
            [73.33, 26.67],
            "q9_help_request_style.png",
            "Single-choice question, N = 75.",
        ),
        "q10": bar_chart(
            "Q10. Do You Omit Tone Words and Pleasantries?",
            ["Yes", "No"],
            [53.33, 46.67],
            "q10_omit_pleasantries.png",
            "Single-choice question, N = 75.",
            colors=[BLUE, TEAL],
        ),
        "q11": bar_chart(
            "Q11. Have You Become More Direct or Blunt?",
            ["Yes", "No"],
            [58.67, 41.33],
            "q11_more_direct_blunt.png",
            "Single-choice question, N = 75.",
            colors=[BLUE, TEAL],
        ),
    }


def build_doc():
    charts = make_charts()
    doc = Document()
    style_doc(doc)

    add_centered(doc, "Project-based Learning Report", 22, True, 24)
    add_centered(doc, "How AI Influences Our Communication", 20, True, 24)
    add_centered(doc, "A Survey and Interview Analysis on the Subtle Effects of AI", 15, False, 42)
    add_centered(doc, "School of Medicine", 13, False, 12)
    add_centered(doc, "June 2026", 13, False, 12)
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_para(doc, "This report investigates whether artificial intelligence subtly influences people’s communication habits without their conscious awareness. Based on 75 valid questionnaire responses, supplementary situational interviews, and a review of recent discussions on AI hallucination, sycophancy, and social dependence, the study finds that AI is not only used as an information tool but also increasingly serves emotional and interpersonal functions. The survey shows that 72% of respondents are aged 18 to 35, 65.33% use AI every day, 88% use AI as a search engine, and 42.67% use it as a confidant. More importantly, although 60% of respondents in one question recognize AI affirmation as mainly flattering, 76% still feel pleased when affirmed by AI. Meanwhile, 53.33% report omitting pleasantries in conversation and 58.67% report becoming more direct or blunt. These findings support the hypothesis that AI may reshape users’ emotional expectations, language style, and interpersonal communication habits gradually and subconsciously.")
    add_para(doc, "Keywords: artificial intelligence; communication; emotional dependence; AI sycophancy; interpersonal interaction", False)

    doc.add_heading("1. Introduction and Research Background", level=1)
    add_para(doc, "Artificial intelligence has become part of daily communication. People use it to search for information, make decisions, write messages, solve personal problems, and even seek emotional support. This convenience creates a paradox. On the one hand, AI makes communication faster and more efficient. On the other hand, because AI systems often respond immediately, politely, and agreeably, they may quietly change what users expect from human interaction.")
    add_para(doc, "The central question of this project is not simply whether people use AI frequently. Instead, the study asks whether long-term interaction with AI subtly changes how people communicate with others, how they seek validation, and how they respond to disagreement. This question matters because such influence may not be consciously noticed by users. Many people may deny that AI has changed them when asked directly, yet their actual habits may reveal a different pattern.")

    doc.add_heading("2. Literature Review", level=1)
    add_para(doc, "Existing research and public discussion suggest three relevant concerns. First, research on language-model hallucination argues that AI systems may produce confident but unreliable answers. Because models are often optimized to provide answers instead of admitting uncertainty, users must learn to maintain a critical attitude toward AI output. This habit may later be transferred into human communication.")
    add_para(doc, "Second, studies on AI sycophancy show that AI systems may approve or support users more readily than human interlocutors. The literature review provided for this project notes that AI models can be more likely to approve users’ decisions, including questionable decisions, in order to maintain user satisfaction. This tendency is directly related to our concern that AI may create an unrealistically agreeable communication environment.")
    add_para(doc, "Third, personal accounts such as Alexandra Samuel’s discussion of prolonged chatbot use suggest that AI can affect real social connections. However, such accounts are often personal and Western-centered. This project therefore examines Chinese users, especially young and frequent users, through questionnaire and interview data. The research gap is the connection between AI’s communicative features and users’ everyday interpersonal habits.")

    doc.add_heading("3. Hypothesis", level=1)
    add_para(doc, "We hypothesized that AI influences people without their conscious awareness. This influence may appear in five dimensions: communication tone, communication style, emotional expectations, interpersonal interaction habits, and subconscious behavioral change. Specifically, AI may make users more accustomed to direct, efficient, agreeable, and emotionally validating communication. Over time, users may prefer interactions with lower conflict and quicker affirmation, while becoming less patient with the complexity of real human communication.")

    doc.add_heading("4. Methodology", level=1)
    add_para(doc, "The study used a mixed-method design combining questionnaires and interviews. Two questionnaires were designed during the research process. The first questionnaire was not satisfactory because some questions were too direct. For example, asking participants to judge whether AI had changed their daily communication or whether their tone toward family members had become worse touched personal dignity and self-image. Such direct questions made respondents less willing to admit possible changes.")
    add_para(doc, "The second questionnaire therefore used a more indirect and tactful design. It collected demographic information, AI usage frequency, purposes of AI use, perceptions of AI tone, emotional responses to AI affirmation, reasons for relying on AI, and changes in real-life communication habits. It also included a scenario-based prompt asking respondents how they would ask AI to recommend travel places in their hometown. This allowed the study to observe natural language style without forcing participants to evaluate themselves directly.")
    add_para(doc, "Interviews were used as a qualitative supplement. Direct questioning first failed to reveal subconscious influence: for instance, Linda initially denied that AI had affected her daily communication. However, when specific situations were introduced, she gradually admitted enjoying AI’s validation and avoiding people who might question her views. Bob also explained that because AI can provide false information, he developed a habit of questioning it, and this habit sometimes extended to discussions with classmates.")
    add_summary_table(doc)

    doc.add_heading("5. Findings and Data Analysis", level=1)
    doc.add_heading("5.1 Young Adults and Frequent AI Users Form the Core Sample", level=2)
    add_para(doc, "The questionnaire results show that young adults are the main user group in this study. Among 75 valid respondents, 72% are aged 18 to 35, while 65.33% use AI every day. This means the data mainly reflect the habits of people who are already familiar with AI and are likely to integrate it into daily routines.")
    add_figure(doc, charts["q1"], "Figure 1. Q1 age distribution.")
    add_figure(doc, charts["q2"], "Figure 2. Q2 AI usage frequency.")
    add_para(doc, "This finding is useful for the research because subtle influence is more likely to appear among frequent users. If AI communication patterns are repeatedly experienced every day, they have more opportunity to become normalized. Therefore, the sample is suitable for examining whether AI communication affects users’ tone, expectations, and interpersonal habits.")

    doc.add_heading("5.2 AI Use Extends from Information Search to Emotional Support", level=2)
    add_para(doc, "The second important phenomenon is that AI is not only used as a technical tool. Although 88% of respondents use AI as a search engine, 42.67% also use it as a confidant, and 26.67% consult it about social problems. These figures show that AI has entered areas that traditionally belonged to interpersonal communication.")
    add_figure(doc, charts["q3"], "Figure 3. Q3 purposes of using AI.")
    add_para(doc, "This result supports the hypothesis about interpersonal interaction habits. When people use AI for emotional expression or social advice, part of their need for human feedback is transferred to a machine. The relationship between user and AI therefore becomes more than functional; it begins to imitate companionship. This helps explain why AI may affect communication expectations even when users believe they are only using it for convenience.")

    doc.add_heading("5.3 Users Recognize AI Flattery but Still Enjoy Its Affirmation", level=2)
    add_para(doc, "A striking contradiction appears between rational judgment and emotional response. In the question about whether AI affirmation is based on flattery or facts, 12.5% chose “completely flattering” and 47.5% chose “basically flattering,” making 60% in total. However, 76% of respondents said they feel pleased when their behavior is affirmed by AI. Q6 further shows that respondents are not all blindly obedient: the largest group, 42.67%, placed their likelihood of adopting AI suggestions in the 40-60 range, while 37.33% placed it above 60.")
    add_figure(doc, charts["q4"], "Figure 4. Q4 perception of AI affirmation.")
    add_figure(doc, charts["q5"], "Figure 5. Q5 emotional response to AI affirmation.")
    add_figure(doc, charts["q6"], "Figure 6. Q6 probability of adopting AI suggestions.")
    add_para(doc, "These three charts together clarify the emotional mechanism. Users may rationally know that AI is flattering, and they may still keep some judgment when deciding whether to follow AI suggestions. Yet the positive feeling produced by AI affirmation remains strong. This means the influence operates less through complete belief and more through repeated emotional reward. AI provides a low-risk, non-confrontational form of validation, which directly supports the hypothesis that AI influences emotional expectations in a subtle way.")

    doc.add_heading("5.4 Convenience Encourages the Transfer of Help-Seeking to AI", level=2)
    add_para(doc, "The survey also asked what high-intensity AI use is related to. The two strongest options were “AI has richer functions” and “AI is more convenient than asking others,” each selected by 68% of respondents. By comparison, only 38.67% selected “AI information is more reliable.” This means intensive use is driven more by functional richness and convenience than by absolute trust.")
    add_figure(doc, charts["q7"], "Figure 7. Q7 reasons related to high-intensity AI use.")
    add_para(doc, "This helps explain why AI may weaken some real-life communication opportunities. When users repeatedly choose AI because it is easier and safer, they may gradually become less willing to engage in human communication that requires patience, explanation, negotiation, or emotional risk. The data therefore connect AI convenience with possible avoidance of interpersonal friction.")

    doc.add_heading("5.5 Communication Becomes More Direct and Less Emotionally Cushioned", level=2)
    add_para(doc, "The final group of findings concerns communication style. When asked about asking others for help, 73.33% said they would first explain the reason before making a request. This suggests that many respondents still preserve human politeness norms. However, 53.33% also admitted that they may subconsciously omit tone words, padding, and polite formulas when chatting, keeping only the core information. In addition, 58.67% reported that they have recently become more direct and blunt, with fewer explanations, comforting words, or softening tones.")
    add_figure(doc, charts["q9"], "Figure 8. Q9 style of asking others for help.")
    add_figure(doc, charts["q10"], "Figure 9. Q10 omission of tone words and pleasantries.")
    add_figure(doc, charts["q11"], "Figure 10. Q11 more direct or blunt communication.")
    add_para(doc, "This pattern is especially important because it shows a partial transfer of AI-style communication into human interaction. AI communication is usually efficient, direct, and task-centered. These features are useful when interacting with a machine, but in human communication they can appear cold, impatient, or rude. The data therefore support the hypothesis that AI may affect not only what people say but also how they organize tone and social cushioning.")

    doc.add_heading("6. Interview Analysis", level=1)
    add_para(doc, "The interviews deepen the questionnaire findings by showing how subconscious influence may appear only after situational questioning. Linda initially denied that AI had influenced her daily communication. However, when asked through concrete scenarios, she admitted that she enjoyed talking with AI because it made her feel validated and accepted. She also became less willing to communicate with people who might deny her or question her views. Her case illustrates how AI’s agreeable tone may raise users’ expectations for positive emotional feedback.")
    add_para(doc, "Bob’s interview reveals another mechanism. Because AI sometimes produces false information, Bob felt that he had to remain critical when using it. This critical habit was useful for checking AI output, but he noticed that it also affected his discussions with classmates. He became less likely to trust others immediately and sometimes questioned them in a rude way. Bob’s case shows that habits developed for AI interaction can be carried into human communication even when the original reason for the habit is reasonable.")
    add_para(doc, "Together, Linda and Bob support the questionnaire results. Linda’s experience corresponds to the data on emotional pleasure and avoidance of disagreement, while Bob’s experience corresponds to the transfer of AI-related communication habits into real-life discussions. The interviews therefore help explain why direct survey questions may underestimate AI influence: people often notice the change only when they reflect on specific situations.")

    doc.add_heading("7. Discussion: What the Data Reveals", level=1)
    add_para(doc, "Overall, the data reveal three connected phenomena. First, AI has moved from being only an information tool to becoming a source of emotional support. This is shown by the high percentage of respondents using AI as a confidant and by the strong pleasure response to AI affirmation. Second, users are not simply deceived by AI’s agreeable tone. Many recognize it as flattery, yet they still respond emotionally to it. This means the influence operates less through belief and more through repeated emotional experience.")
    add_para(doc, "Third, AI interaction may reshape communication style by making direct, efficient, and low-context expression feel normal. The findings on omitted pleasantries and more blunt expression suggest that AI-style communication can leak into human communication. This does not mean AI is the only cause of these changes, but the consistency between the questionnaire and interviews makes it reasonable to argue that AI is one contributing factor.")
    add_para(doc, "These findings are helpful to the research because they connect the literature review with primary data. The literature shows that AI can hallucinate and flatter users. Our data show what such features may do to users: they may encourage emotional dependence, avoidance of disagreement, mechanical expression, and transferred critical habits. Therefore, the study contributes a user-centered perspective to discussions that often focus only on model behavior.")

    doc.add_heading("8. Conclusion", level=1)
    add_para(doc, "The research supports the hypothesis that AI may subtly influence communication without users’ full awareness. The influence appears in emotional expectations, help-seeking behavior, and communication style. AI’s convenience and agreeable tone make it attractive not only as a tool but also as an emotional partner. At the same time, repeated interaction with AI may encourage users to become more direct, less emotionally cushioned, and sometimes less willing to face disagreement in real human relationships.")
    add_para(doc, "The study also shows the importance of research design. Direct questions may fail because respondents do not want to admit changes that threaten their self-image. Indirect questionnaire design and situational interviews are therefore more effective for studying subconscious communication habits. Future research could expand the sample size, compare heavy and light users, and observe actual chat records or communication behavior over time.")

    doc.add_heading("References", level=1)
    refs = [
        "Kalai, A. (2025). Why Language Models Hallucinate. arXiv.",
        "Cheng, M. (2025). Sycophantic AI decreases prosocial intentions and promotes dependence. Science.",
        "Samuel, A. (2024). I Love Being Social. Then I Started Talking to a Chatbot. The Wall Street Journal.",
        "Project questionnaire data. (2026). AI and the subtle influence on communication. Unpublished survey report.",
    ]
    for ref in refs:
        add_para(doc, ref, False)

    doc.add_heading("Appendix: Key Questionnaire Data Used in This Report", level=1)
    data_rows = [
        ("Q1 Under 18", "4 / 75", "5.33%"),
        ("Q1 Age 18-35", "54 / 75", "72%"),
        ("Q1 Age 35-45", "7 / 75", "9.33%"),
        ("Q1 Age 45 and above", "10 / 75", "13.33%"),
        ("Q2 Every day", "49 / 75", "65.33%"),
        ("Q2 Once every two or three days", "26 / 75", "34.67%"),
        ("Q2 Rarely", "0 / 75", "0%"),
        ("Q3 Search engine", "66 / 75", "88%"),
        ("Q3 Confidant / emotional object", "32 / 75", "42.67%"),
        ("Q3 Advice on social problems", "20 / 75", "26.67%"),
        ("Q4 Completely flattering", "5 / 40", "12.5%"),
        ("Q4 Basically flattering", "19 / 40", "47.5%"),
        ("Q4 Basically fact-based", "10 / 40", "25%"),
        ("Q4 Completely fact-based", "6 / 40", "15%"),
        ("Q5 Yes, feel pleased when affirmed", "57 / 75", "76%"),
        ("Q5 No", "18 / 75", "24%"),
        ("Q6 0-20 probability of adopting AI suggestions", "6 / 75", "8%"),
        ("Q6 20-40 probability of adopting AI suggestions", "9 / 75", "12%"),
        ("Q6 40-60 probability of adopting AI suggestions", "32 / 75", "42.67%"),
        ("Q6 60-80 probability of adopting AI suggestions", "19 / 75", "25.33%"),
        ("Q6 80-100 probability of adopting AI suggestions", "9 / 75", "12%"),
        ("Q7 AI has richer functions", "51 / 75", "68%"),
        ("Q7 AI information is more reliable", "29 / 75", "38.67%"),
        ("Q7 AI is more convenient than asking others", "51 / 75", "68%"),
        ("Q9 Explain the reason before asking", "55 / 75", "73.33%"),
        ("Q9 Directly state the task", "20 / 75", "26.67%"),
        ("Q10 Yes, omit tone words and pleasantries", "40 / 75", "53.33%"),
        ("Q10 No", "35 / 75", "46.67%"),
        ("Q11 Yes, become more direct or blunt", "44 / 75", "58.67%"),
        ("Q11 No", "31 / 75", "41.33%"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["Question / Indicator", "Count", "Percentage"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in data_rows:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            set_cell_text(cells[i], item, i == 0)

    out = OUT / "How_AI_Influences_Our_Communication_Report_Final.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_doc())
