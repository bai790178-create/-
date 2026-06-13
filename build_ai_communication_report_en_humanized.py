from pathlib import Path

import build_ai_communication_report as base
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("report_output")
CHARTS = OUT / "charts_humanized_en"
OUT.mkdir(exist_ok=True)
CHARTS.mkdir(exist_ok=True)
base.CHARTS = CHARTS


def clear_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)


def add_figure_pair(doc, left_item, right_item):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(table)
    for col, item in enumerate([left_item, right_item]):
        image_path, caption = item
        image_cell = table.rows[0].cells[col]
        cap_cell = table.rows[1].cells[col]
        image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cap_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = image_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.add_run().add_picture(str(image_path), width=Inches(3.05))
        cp = cap_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = None
        cp.paragraph_format.line_spacing = 1.05
        r = cp.add_run(caption)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(9.5)
        r.italic = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = None
    spacer.paragraph_format.space_after = Pt(6)


def make_charts():
    return base.make_charts()


def add_appendix(doc):
    rows = [
        ("Q1 Under 18", "4 / 75", "5.33%"),
        ("Q1 Age 18-35", "54 / 75", "72%"),
        ("Q1 Age 35-45", "7 / 75", "9.33%"),
        ("Q1 Age 45 and above", "10 / 75", "13.33%"),
        ("Q2 Every day", "49 / 75", "65.33%"),
        ("Q2 Once every two or three days", "26 / 75", "34.67%"),
        ("Q2 Rarely", "0 / 75", "0%"),
        ("Q3 Search engine", "66 / 75", "88%"),
        ("Q3 Confidant / emotional object", "32 / 75", "42.67%"),
        ("Q3 Advice on social problems", "20 / 75", "26.67%"),
        ("Q4 Completely flattering", "5 / 40", "12.5%"),
        ("Q4 Basically flattering", "19 / 40", "47.5%"),
        ("Q4 Basically fact-based", "10 / 40", "25%"),
        ("Q4 Completely fact-based", "6 / 40", "15%"),
        ("Q5 Yes, feel pleased when affirmed", "57 / 75", "76%"),
        ("Q5 No", "18 / 75", "24%"),
        ("Q6 0-20 probability of adopting AI suggestions", "6 / 75", "8%"),
        ("Q6 20-40 probability of adopting AI suggestions", "9 / 75", "12%"),
        ("Q6 40-60 probability of adopting AI suggestions", "32 / 75", "42.67%"),
        ("Q6 60-80 probability of adopting AI suggestions", "19 / 75", "25.33%"),
        ("Q6 80-100 probability of adopting AI suggestions", "9 / 75", "12%"),
        ("Q7 AI has richer functions", "51 / 75", "68%"),
        ("Q7 AI information is more reliable", "29 / 75", "38.67%"),
        ("Q7 AI is more convenient than asking others", "51 / 75", "68%"),
        ("Q9 Explain the reason before asking", "55 / 75", "73.33%"),
        ("Q9 Directly state the task", "20 / 75", "26.67%"),
        ("Q10 Yes, omit tone words and pleasantries", "40 / 75", "53.33%"),
        ("Q10 No", "35 / 75", "46.67%"),
        ("Q11 Yes, become more direct or blunt", "44 / 75", "58.67%"),
        ("Q11 No", "31 / 75", "41.33%"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["Question / Indicator", "Count", "Percentage"]):
        base.set_cell_text(table.rows[0].cells[i], h, True)
        base.set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            base.set_cell_text(cells[i], value, i == 0)


def build_doc():
    charts = make_charts()
    doc = Document()
    base.style_doc(doc)

    base.add_centered(doc, "Project-based Learning Report", 22, True, 24)
    base.add_centered(doc, "How AI Influences Our Communication", 20, True, 24)
    base.add_centered(doc, "A Survey and Interview Report on Subtle Changes in Everyday Interaction", 15, False, 42)
    base.add_centered(doc, "School of Medicine", 13, False, 12)
    base.add_centered(doc, "June 2026", 13, False, 12)
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    base.add_para(doc, "This report looks at a question that is easy to overlook: whether AI changes the way people communicate even when they do not feel changed. The study uses 75 valid questionnaire responses and several student interviews. The data suggest that AI is no longer used only for quick answers. It is also becoming a place for reassurance, advice, and low-pressure conversation. In the survey, 72% of respondents were aged 18 to 35, 65.33% used AI every day, 88% used AI as a search engine, and 42.67% used it as a confidant. A more interesting tension appears in the emotional data: although 60% of respondents in Q4 viewed AI affirmation as mainly flattering, 76% still felt pleased when AI affirmed them. At the same time, 53.33% reported omitting pleasantries in conversation, and 58.67% felt that their recent communication had become more direct. These results do not prove that AI is the only cause, but they do show a clear pattern worth taking seriously.")
    base.add_para(doc, "Keywords: artificial intelligence; communication; emotional validation; AI sycophancy; interpersonal interaction", False)

    doc.add_heading("1. Introduction and Research Background", level=1)
    base.add_para(doc, "AI has slipped into ordinary communication. People ask it for information, help with writing, social advice, and sometimes emotional comfort. That convenience is useful, but it also creates a quieter problem. AI usually answers quickly, patiently, and agreeably. Human conversation is not always like that. It includes misunderstanding, waiting, rejection, disagreement, and the need to soften one’s tone.")
    base.add_para(doc, "Our project therefore does not ask only whether people use AI. It asks whether repeated contact with AI may slowly reshape what people expect from communication. The change may be hard to notice because users often do not describe themselves as being influenced by AI. They may only recognize the influence when they are asked about specific situations.")

    doc.add_heading("2. Literature Review", level=1)
    base.add_para(doc, "The literature we reviewed points to three concerns. First, discussions of language-model hallucination show that AI can produce confident but unreliable answers. This pushes users to stay critical when they interact with AI. Second, research on AI sycophancy suggests that AI may support or approve users more readily than people do. This matters because approval can feel emotionally rewarding even when users know it is not fully objective.")
    base.add_para(doc, "Third, personal accounts of chatbot use suggest that prolonged interaction with AI may affect real relationships. However, those accounts are often individual cases and are usually written in Western contexts. Our project tries to add a small user-centered perspective from Chinese students and young users, a group that tends to be familiar with AI and generally open to using it.")

    doc.add_heading("3. Hypothesis", level=1)
    base.add_para(doc, "We hypothesized that AI influences people in ways they may not consciously notice. The influence may appear in communication tone, expression style, emotional expectations, and real-life interaction habits. In simple terms, people who often talk with AI may become more used to direct, efficient, agreeable, and validating communication. This may make ordinary human communication feel slower, less comfortable, or more emotionally demanding.")

    doc.add_heading("4. Methodology", level=1)
    base.add_para(doc, "The project combined questionnaire data with interviews. We first designed a direct questionnaire, but the results were not very useful. Some questions asked participants to judge whether AI had changed them or whether their tone toward family members had become worse. Those questions were too close to personal self-image, so respondents had reasons to answer defensively.")
    base.add_para(doc, "The second questionnaire was more indirect. It asked about age, AI use frequency, purposes of AI use, perceptions of AI affirmation, emotional responses, reasons for high-intensity AI use, and recent communication habits. It also included a scenario about asking AI to recommend places in one’s hometown. That kind of question helped us observe language style without directly accusing anyone of changing.")
    base.add_para(doc, "The interviews followed the same logic. Linda first said AI had not influenced her, but when we moved to concrete situations, she admitted that AI’s validation made her feel accepted and that she sometimes avoided people who might challenge her. Bob described a different pattern: because AI can be wrong, he learned to question it often, but that questioning habit sometimes carried over into discussions with classmates.")
    base.add_summary_table(doc)

    doc.add_heading("5. Findings and Data Analysis", level=1)
    doc.add_heading("5.1 Young and Frequent Users Dominate the Sample", level=2)
    base.add_para(doc, "The sample is strongly concentrated among young users. In Q1, 72% of respondents were aged 18 to 35. Q2 shows that 65.33% used AI every day, while the rest used it at least once every two or three days. No respondent selected “rarely.” This matters because the research question is about subtle influence, and subtle influence is more likely to appear when AI is used repeatedly.")
    add_figure_pair(doc, (charts["q1"], "Figure 1. Q1 age distribution."), (charts["q2"], "Figure 2. Q2 AI usage frequency."))
    base.add_para(doc, "Because most respondents are regular AI users, their answers are useful for observing habits rather than one-time impressions. The data do not represent every age group equally, but they fit the group most likely to experience AI as part of daily communication.")

    doc.add_heading("5.2 AI Is Used for More Than Information", level=2)
    base.add_para(doc, "Q3 shows a shift from tool use to emotional use. The largest group, 88%, used AI as a search engine. That result is expected. What is more revealing is that 42.67% also used AI as a confidant, and 26.67% used it to ask about social problems. These numbers suggest that AI is entering spaces that used to depend more on friends, classmates, or family members.")
    base.add_figure(doc, charts["q3"], "Figure 3. Q3 purposes of using AI.")
    base.add_para(doc, "This helps our study because it shows where influence may happen. If AI were only a calculator or a dictionary, its effect on communication would probably be limited. But when it becomes a place for emotional expression and social advice, it begins to shape what users expect from a listener.")

    doc.add_heading("5.3 Users Know AI May Flatter Them, but the Feeling Still Works", level=2)
    base.add_para(doc, "The strongest pattern in the data is the gap between what users know and what they feel. In Q4, 60% of respondents said AI’s affirmation was completely or basically flattering. Yet Q5 shows that 76% still felt pleased when AI affirmed them. Q6 adds one more layer: most respondents were not blindly obedient, since the largest group placed their probability of adopting AI suggestions in the 40-60 range.")
    add_figure_pair(doc, (charts["q4"], "Figure 4. Q4 perception of AI affirmation."), (charts["q5"], "Figure 5. Q5 emotional response to AI affirmation."))
    base.add_figure(doc, charts["q6"], "Figure 6. Q6 probability of adopting AI suggestions.")
    base.add_para(doc, "This is why the influence is subtle. Users may remain rational enough to doubt AI, but the emotional reward still works. AI offers validation without the risk of embarrassment or disagreement. Linda’s interview supports this point: she knew AI was not the same as a person, but she still enjoyed the feeling of being accepted.")

    doc.add_heading("5.4 Convenience Pushes Help-Seeking Toward AI", level=2)
    base.add_para(doc, "Q7 suggests that frequent AI use is not mainly about blind trust. The two strongest reasons were richer functions and convenience, both at 68%. Only 38.67% selected higher reliability. In other words, people may use AI heavily because it is easy to access, not because they believe it is always correct.")
    base.add_figure(doc, charts["q7"], "Figure 7. Q7 reasons related to high-intensity AI use.")
    base.add_para(doc, "This is important for communication. Asking AI has almost no social cost. Users do not need to worry about interrupting someone, being judged, or receiving a negative reaction. Over time, that low-friction habit may make real human communication feel less attractive.")

    doc.add_heading("5.5 Human Communication Becomes More Direct", level=2)
    base.add_para(doc, "The final group of questions connects AI use with real communication style. Q9 shows that 73.33% still prefer to explain the reason before asking others for help, so politeness norms have not disappeared. However, Q10 shows that 53.33% sometimes omit tone words, padding, and pleasantries, while Q11 shows that 58.67% feel they have become more direct or blunt in recent communication.")
    add_figure_pair(doc, (charts["q9"], "Figure 8. Q9 style of asking others for help."), (charts["q10"], "Figure 9. Q10 omission of tone words and pleasantries."))
    base.add_figure(doc, charts["q11"], "Figure 10. Q11 more direct or blunt communication.")
    base.add_para(doc, "This does not mean AI has erased social manners. The picture is more mixed. People still know how human communication should work, but some AI-style habits may be entering daily expression: shorter wording, fewer softeners, and a stronger focus on the core task.")

    doc.add_heading("6. Interview Analysis", level=1)
    base.add_para(doc, "The interviews make the numbers easier to understand. Linda’s case shows the emotional side of the influence. She did not first describe herself as changed by AI. Only after discussing specific situations did she say that AI’s validation felt comfortable and that disagreement from real people felt harder to face. Her response matches the survey result that many users enjoy AI affirmation even when they know it may be flattering.")
    base.add_para(doc, "Bob’s case shows a different kind of transfer. His critical attitude toward AI was reasonable because AI can be wrong. The problem was that this habit did not always stay inside AI use. In conversations with classmates, he sometimes became more suspicious and questioning than he intended. This supports our idea that habits built around AI can leak into human interaction.")

    doc.add_heading("7. Discussion: What the Data Reveal", level=1)
    base.add_para(doc, "The data reveal three main phenomena. First, AI is becoming a place for both information and emotional support. Second, users can recognize AI flattery and still respond positively to it. Third, the directness of AI interaction may influence the tone of human communication.")
    base.add_para(doc, "These findings connect the literature review with our own data. The literature explains why AI may hallucinate or flatter users. Our survey and interviews show what that may mean for users: more reliance on low-friction validation, more avoidance of uncomfortable disagreement, and a gradual shift toward concise but less emotionally cushioned expression.")

    doc.add_heading("8. Conclusion", level=1)
    base.add_para(doc, "The study supports the hypothesis that AI can subtly influence communication. The influence is not dramatic or immediately visible. It appears through repeated small experiences: being answered instantly, being affirmed politely, avoiding the awkwardness of asking real people, and getting used to direct task-centered language.")
    base.add_para(doc, "The project also shows why indirect research methods matter. If we ask people directly whether AI has changed them, many will probably say no. But when questions focus on specific habits and situations, the pattern becomes clearer. Future research could compare heavy and light AI users, expand the sample, and examine real communication records over time.")

    doc.add_heading("References", level=1)
    for ref in [
        "Kalai, A. (2025). Why Language Models Hallucinate. arXiv.",
        "Cheng, M. (2025). Sycophantic AI decreases prosocial intentions and promotes dependence. Science.",
        "Samuel, A. (2024). I Love Being Social. Then I Started Talking to a Chatbot. The Wall Street Journal.",
        "Project questionnaire data. (2026). AI and the subtle influence on communication. Unpublished survey report.",
    ]:
        base.add_para(doc, ref, False)

    doc.add_heading("Appendix: Key Questionnaire Data Used in This Report", level=1)
    add_appendix(doc)

    out = OUT / "How_AI_Influences_Our_Communication_Report_Humanized.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_doc())
